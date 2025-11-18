# """
# tools.py

# Agent Tools:
# - think
# - create_docx_report
# - send_email
# """

# import os
# from pathlib import Path
# from typing import Optional, Any, Dict

# from dotenv import load_dotenv
# load_dotenv()

# from pydantic import BaseModel, Field

# # LangChain tool decorator (all LC 1.0.x)
# try:
#     from langchain_core.tools import tool
# except ImportError:
#     from langchain.tools import tool

# # python-docx
# try:
#     from docx import Document
# except Exception:
#     Document = None

# import smtplib
# from email.message import EmailMessage


# EMAIL_USER = os.getenv("EMAIL_USER")
# EMAIL_PASS = os.getenv("EMAIL_PASS")

# REPORTS_DIR = Path("reports")
# REPORTS_DIR.mkdir(exist_ok=True)


# # =======================================================
# # THINK TOOL
# # =======================================================
# class ThinkInput(BaseModel):
#     note: str = Field(..., description="Internal planning/scratchpad note")

# @tool(args_schema=ThinkInput)
# def think(input: ThinkInput) -> Dict[str, Any]:
#     """
#     A side-effect-free scratchpad used by the planner/critic to store
#     structured reasoning notes inside the agent's state.
#     """
#     return {"status": "ok", "note": input.note}


# # =======================================================
# # CREATE DOCX REPORT
# # =======================================================
# class CreateReportInput(BaseModel):
#     title: str
#     content: str

# @tool(args_schema=CreateReportInput)
# def create_docx_report(input: CreateReportInput) -> Dict[str, Any]:
#     """
#     Creates a DOCX report file containing a title and body text.
#     Returns structured metadata describing the file generated.
#     """
#     if Document is None:
#         return {"status": "error", "error": "python-docx not installed"}

#     filename = f"{input.title.strip().replace(' ', '_')}.docx"
#     filepath = REPORTS_DIR / filename

#     try:
#         doc = Document()
#         doc.add_heading(input.title, level=1)
#         doc.add_paragraph(input.content)
#         doc.save(filepath)
#         return {"status": "ok", "path": str(filepath), "filename": filename}
#     except Exception as e:
#         return {"status": "error", "error": str(e)}


# # =======================================================
# # SEND EMAIL
# # =======================================================
# class SendEmailInput(BaseModel):
#     recipient: str
#     subject: str
#     body: str
#     attachment_path: Optional[str] = None

# @tool(args_schema=SendEmailInput)
# def send_email(input: SendEmailInput) -> Dict[str, Any]:
#     """
#     Sends an email using Gmail SMTP + optional attachment.
#     Requires EMAIL_USER and EMAIL_PASS in .env
#     Returns structured metadata describing success or failure.
#     """
#     if not EMAIL_USER or not EMAIL_PASS:
#         return {"status": "error", "error": "Missing EMAIL_USER/EMAIL_PASS"}

#     msg = EmailMessage()
#     msg["From"] = EMAIL_USER
#     msg["To"] = input.recipient
#     msg["Subject"] = input.subject
#     msg.set_content(input.body)

#     if input.attachment_path:
#         p = Path(input.attachment_path)
#         if not p.exists():
#             return {"status": "error", "error": f"Attachment not found: {p}"}

#         try:
#             with open(p, "rb") as f:
#                 data = f.read()
#             msg.add_attachment(
#                 data,
#                 maintype="application",
#                 subtype="octet-stream",
#                 filename=p.name,
#             )
#         except Exception as e:
#             return {"status": "error", "error": f"Attachment read failed: {e}"}

#     try:
#         with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
#             smtp.login(EMAIL_USER, EMAIL_PASS)
#             smtp.send_message(msg)

#         return {"status": "ok", "message": "Email sent"}

#     except Exception as e:
#         return {"status": "error", "error": str(e)}


# # =======================================================
# # EXPORT TOOL LIST
# # =======================================================
# TOOLS = [think, create_docx_report, send_email]



# """
# tools.py

# Agent Tools:
# - think
# - create_docx_report
# - send_email

# This file exports TOOLS = [tool_obj, ...] where each tool_obj has:
#  - .name  (string)
#  - .func  (callable that accepts a single argument)
# """

# import os
# from pathlib import Path
# from typing import Optional, Any, Dict, Callable
# from dotenv import load_dotenv
# load_dotenv()

# from pydantic import BaseModel, Field

# # Try imports for LangChain's tool decorator (support multiple versions)
# try:
#     from langchain_core.tools import tool
# except Exception:
#     try:
#         from langchain.tools import tool
#     except Exception:
#         tool = None  # we'll handle absence below

# # python-docx optional
# try:
#     from docx import Document
# except Exception:
#     Document = None

# import smtplib
# from email.message import EmailMessage

# EMAIL_USER = os.getenv("EMAIL_USER")
# EMAIL_PASS = os.getenv("EMAIL_PASS")

# REPORTS_DIR = Path("reports")
# REPORTS_DIR.mkdir(exist_ok=True)


# # ---------------------------
# # Tool: THINK (scratchpad)
# # ---------------------------
# class ThinkInput(BaseModel):
#     note: str = Field(..., description="Internal planning/scratchpad note")

# # If decorator available, use it; otherwise keep plain function and wrap later
# if tool is not None:
#     @tool(args_schema=ThinkInput)
#     def think(input: ThinkInput) -> Dict[str, Any]:
#         """
#         A side-effect-free scratchpad. Returns {"status":"ok","note":...}
#         """
#         return {"status": "ok", "note": input.note}
# else:
#     def think(input: ThinkInput) -> Dict[str, Any]:
#         """
#         A side-effect-free scratchpad. Returns {"status":"ok","note":...}
#         """
#         return {"status": "ok", "note": input.note}


# # ---------------------------
# # Tool: CREATE DOCX REPORT
# # ---------------------------
# class CreateReportInput(BaseModel):
#     title: str = Field(..., description="Report title")
#     content: str = Field(..., description="Report body text")

# if tool is not None:
#     @tool(args_schema=CreateReportInput)
#     def create_docx_report(input: CreateReportInput) -> Dict[str, Any]:
#         """
#         Create a DOCX report. Returns {"status":"ok","path":..., "filename":...}
#         """
#         if Document is None:
#             return {"status": "error", "error": "python-docx not installed"}
#         filename = f"{input.title.strip().replace(' ', '_')}.docx"
#         filepath = REPORTS_DIR / filename
#         try:
#             doc = Document()
#             doc.add_heading(input.title, level=1)
#             doc.add_paragraph(input.content)
#             doc.save(filepath)
#             return {"status": "ok", "path": str(filepath), "filename": filename}
#         except Exception as e:
#             return {"status": "error", "error": str(e)}
# else:
#     def create_docx_report(input: CreateReportInput) -> Dict[str, Any]:
#         if Document is None:
#             return {"status": "error", "error": "python-docx not installed"}
#         filename = f"{input.title.strip().replace(' ', '_')}.docx"
#         filepath = REPORTS_DIR / filename
#         try:
#             doc = Document()
#             doc.add_heading(input.title, level=1)
#             doc.add_paragraph(input.content)
#             doc.save(filepath)
#             return {"status": "ok", "path": str(filepath), "filename": filename}
#         except Exception as e:
#             return {"status": "error", "error": str(e)}


# # ---------------------------
# # Tool: SEND EMAIL
# # ---------------------------
# class SendEmailInput(BaseModel):
#     recipient: str = Field(..., description="Recipient address")
#     subject: str = Field(..., description="Subject")
#     body: str = Field(..., description="Plain text body")
#     attachment_path: Optional[str] = Field(None, description="Optional file path")

# if tool is not None:
#     @tool(args_schema=SendEmailInput)
#     def send_email(input: SendEmailInput) -> Dict[str, Any]:
#         """
#         Send email via Gmail SMTP. Returns structured status dict.
#         """
#         if not EMAIL_USER or not EMAIL_PASS:
#             return {"status": "error", "error": "EMAIL_USER/EMAIL_PASS not set"}
#         msg = EmailMessage()
#         msg["From"] = EMAIL_USER
#         msg["To"] = input.recipient
#         msg["Subject"] = input.subject
#         msg.set_content(input.body)

#         if input.attachment_path:
#             p = Path(input.attachment_path)
#             if not p.exists():
#                 return {"status": "error", "error": f"Attachment not found: {p}"}
#             try:
#                 with open(p, "rb") as f:
#                     data = f.read()
#                 msg.add_attachment(data, maintype="application", subtype="octet-stream", filename=p.name)
#             except Exception as e:
#                 return {"status": "error", "error": f"Attachment read failed: {e}"}

#         try:
#             with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
#                 smtp.login(EMAIL_USER, EMAIL_PASS)
#                 smtp.send_message(msg)
#             return {"status": "ok", "message": "Email sent"}
#         except Exception as e:
#             return {"status": "error", "error": str(e)}
# else:
#     def send_email(input: SendEmailInput) -> Dict[str, Any]:
#         if not EMAIL_USER or not EMAIL_PASS:
#             return {"status": "error", "error": "EMAIL_USER/EMAIL_PASS not set"}
#         msg = EmailMessage()
#         msg["From"] = EMAIL_USER
#         msg["To"] = input.recipient
#         msg["Subject"] = input.subject
#         msg.set_content(input.body)
#         if input.attachment_path:
#             p = Path(input.attachment_path)
#             if not p.exists():
#                 return {"status": "error", "error": f"Attachment not found: {p}"}
#             try:
#                 with open(p, "rb") as f:
#                     data = f.read()
#                 msg.add_attachment(data, maintype="application", subtype="octet-stream", filename=p.name)
#             except Exception as e:
#                 return {"status": "error", "error": f"Attachment read failed: {e}"}
#         try:
#             with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
#                 smtp.login(EMAIL_USER, EMAIL_PASS)
#                 smtp.send_message(msg)
#             return {"status": "ok", "message": "Email sent"}
#         except Exception as e:
#             return {"status": "error", "error": str(e)}


# # ---------------------------
# # Export TOOLS as uniform objects
# # ---------------------------
# class _SimpleTool:
#     def __init__(self, name: str, func: Callable[[Any], Any]):
#         self.name = name
#         self.func = func


# def _make_callable_wrapper(obj):
#     """
#     Accept decorated tool object (StructuredTool) or plain function.
#     Return a callable that takes either:
#       - a dict (kwargs for pydantic) or
#       - a single string (mapped to common param)
#     The wrapper returns a Python-native dict/result.
#     """
#     # If it's a LangChain StructuredTool-like object
#     if hasattr(obj, "name") and (hasattr(obj, "func") or hasattr(obj, "__call__")):
#         tool_name = getattr(obj, "name")
#         # If it exposes `.func`, use that; else call the object itself.
#         if hasattr(obj, "func") and callable(getattr(obj, "func")):
#             underlying = obj.func
#         else:
#             underlying = obj  # callable object

#         def wrapper(arg):
#             # if arg already looks like a dict, pass through
#             try:
#                 if arg is None:
#                     # some tools expect pydantic model - try passing empty dict
#                     return underlying({})
#                 if isinstance(arg, dict):
#                     return underlying(arg)
#                 # if it's a pydantic model instance or has .dict()
#                 if hasattr(arg, "dict"):
#                     return underlying(arg)
#                 # otherwise pass raw value (string)
#                 return underlying(arg)
#             except Exception as e:
#                 return {"status": "error", "error": f"Tool call failed: {e}"}

#         return tool_name, wrapper

#     # Fallback: plain python function
#     if callable(obj):
#         tool_name = getattr(obj, "__name__", "tool")
#         def wrapper(arg):
#             try:
#                 # If pydantic model expected, try passing a dict
#                 if arg is None:
#                     return obj({})
#                 if isinstance(arg, dict):
#                     return obj(**arg) if hasattr(obj, "__call__") else obj(arg)
#                 if hasattr(arg, "dict"):
#                     return obj(arg)
#                 # string -> try pass as single positional arg
#                 return obj(arg)
#             except TypeError:
#                 # fallback to passing whole dict as single arg
#                 try:
#                     return obj(arg)
#                 except Exception as e:
#                     return {"status": "error", "error": f"Tool call failed: {e}"}
#             except Exception as e:
#                 return {"status": "error", "error": f"Tool error: {e}"}
#         return tool_name, wrapper

#     # Unknown object
#     return None, None


# # Build uniform TOOLS list
# _raw_tools = [think, create_docx_report, send_email]
# TOOLS = []
# for t in _raw_tools:
#     name, call = _make_callable_wrapper(t)
#     if name and call:
#         TOOLS.append(_SimpleTool(name, call))

# # For backward-compat convenience offer a name->tool mapping
# TOOLS_BY_NAME = {t.name.lower(): t for t in TOOLS}


"""
tools.py

Agent Tools:
- think
- create_docx_report
- send_email

This file exports TOOLS = [tool_obj, ...] and TOOLS_BY_NAME = {name: tool_obj}
where each tool_obj has:
 - .name  (string)
 - .func  (a callable that accepts a dictionary)
"""
# tools.py (Full Corrected Code with Pandoc Integration)

import os
import tempfile
from pathlib import Path
from typing import Optional, Any, Dict, Callable, List
from dotenv import load_dotenv
from pydantic import BaseModel, Field

# --- NEW: Import pypandoc ---
try:
    import pypandoc
    PANDOC_AVAILABLE = True
except ImportError:
    PANDOC_AVAILABLE = False

# python-docx is now a fallback, but we can keep the import
try:
    from docx import Document
except Exception:
    Document = None

import smtplib
from email.message import EmailMessage

load_dotenv()

EMAIL_USER = os.getenv("EMAIL_USER")
EMAIL_PASSWORD = os.getenv("EMAIL_PASSWORD") # Corrected from your last question

REPORTS_DIR = Path("reports")
REPORTS_DIR.mkdir(exist_ok=True)


# --- Tool Input Schemas (no changes here) ---
class ThinkInput(BaseModel):
    note: str = Field(..., description="Internal planning or reasoning note.")

class CreateReportInput(BaseModel):
    title: str = Field(..., description="The title of the report.")
    content: str = Field(..., description="The main body content of the report, written in Markdown.")

class SendEmailInput(BaseModel):
    recipient: str = Field(..., description="The recipient's email address.")
    subject: str = Field(..., description="The subject of the email.")
    body: str = Field(..., description="The plain text body of the email.")
    attachment_path: Optional[str] = Field(None, description="Optional local path to a file to attach.")


# --- Tool Functions ---
def think(input: ThinkInput) -> Dict[str, Any]:
    """A scratchpad for the agent to write down its thoughts."""
    return {"status": "ok", "note": input.note}

# --- START: REPLACED FUNCTION ---
def create_docx_report(input: CreateReportInput) -> Dict[str, Any]:
    """
    Creates a DOCX report file from Markdown content using Pandoc.
    """
    if not PANDOC_AVAILABLE:
        return {"status": "error", "error": "The 'pypandoc' library is not installed. Please run 'pip install pypandoc'."}

    filename = f"{input.title.strip().replace(' ', '_')}.docx"
    filepath = REPORTS_DIR / filename

    try:
        # Pandoc works best by converting files, so we write the content to a temporary file
        with tempfile.NamedTemporaryFile(mode='w+', delete=False, suffix='.md', encoding='utf-8') as temp_md:
            # Add the title as a top-level Markdown header
            temp_md.write(f"# {input.title}\n\n")
            temp_md.write(input.content)
            temp_md_path = temp_md.name

        # Use pypandoc to convert the Markdown file to a DOCX file
        pypandoc.convert_file(temp_md_path, 'docx', outputfile=str(filepath))

        return {"status": "ok", "path": str(filepath), "filename": filename}

    except Exception as e:
        # This will catch errors if Pandoc is not installed correctly
        return {"status": "error", "error": f"Failed to create DOCX with Pandoc: {e}. Is Pandoc installed on your system?"}
    
    finally:
        # Clean up the temporary file
        if 'temp_md_path' in locals() and os.path.exists(temp_md_path):
            os.remove(temp_md_path)
# --- END: REPLACED FUNCTION ---


def send_email(input: SendEmailInput) -> Dict[str, Any]:
    """Sends an email using Gmail SMTP."""
    if not EMAIL_USER or not EMAIL_PASSWORD:
        return {"status": "error", "error": "Email credentials (EMAIL_USER, EMAIL_PASSWORD) are not set in the environment."}
    
    msg = EmailMessage()
    msg["From"] = EMAIL_USER
    msg["To"] = input.recipient
    msg["Subject"] = input.subject
    msg.set_content(input.body)

    if input.attachment_path:
        p = Path(input.attachment_path)
        if not p.exists():
            return {"status": "error", "error": f"Attachment file not found at: {p}"}
        try:
            with open(p, "rb") as f:
                data = f.read()
            msg.add_attachment(data, maintype="application", subtype="octet-stream", filename=p.name)
        except Exception as e:
            return {"status": "error", "error": f"Failed to read attachment: {e}"}

    try:
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
            smtp.login(EMAIL_USER, EMAIL_PASSWORD)
            smtp.send_message(msg)
        return {"status": "ok", "message": "Email sent successfully."}
    except Exception as e:
        return {"status": "error", "error": f"Failed to send email: {e}"}

# --- Uniform Tool Export (no changes here) ---
class SimpleTool:
    def __init__(self, name: str, func: Callable, args_schema: BaseModel):
        self.name = name
        self._raw_func = func
        self.args_schema = args_schema

    def func(self, arg: Dict[str, Any]):
        try:
            parsed = self.args_schema(**arg) if isinstance(arg, dict) else arg
            return self._raw_func(parsed)
        except Exception as e:
            return {"status": "error", "error": f"Tool execution failed: {e}"}

TOOLS: List[SimpleTool] = [
    SimpleTool(name="think", func=think, args_schema=ThinkInput),
    SimpleTool(name="create_docx_report", func=create_docx_report, args_schema=CreateReportInput),
    SimpleTool(name="send_email", func=send_email, args_schema=SendEmailInput),
]

TOOLS_BY_NAME: Dict[str, SimpleTool] = {t.name.lower(): t for t in TOOLS}